#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Insert in-text citation markers ([n]) into the (manually-edited) report.

Operates IN PLACE on docs/深度学习大作业_邹研泽_修订.docx without rebuilding, so
the user's manual changes (Word equations etc.) are preserved. Markers are
inserted at the run level (no run flattening), so existing fonts / Word formulas
are untouched. The reference list [1]-[13] already exists in the document, so it
is not modified here.
Run: E:/anaconda3/envs/dl/python.exe scripts/add_citations.py
"""
import copy
import re
import shutil
import sys
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

DOCX = "docs/深度学习大作业_邹研泽_修订.docx"

# (paragraph locator, [(anchor, marker_inserted_right_after_anchor), ...])
# Order within a paragraph is left-to-right; later anchors may rely on earlier
# inserts (e.g. the Qwen2.5-Math sentence is placed after "...行为[5][6]。").
OPS = [
    # ---- 1. 引言 ----
    ("推动了语言模型推理研究的发展",
     [("语言模型推理研究的发展", "[1]")]),
    ("DeepSeek-R1 和 DeepSeekMath 等工作显示",
     [("推动模型形成更强的推理行为", "[5][6]"),
      ("推动模型形成更强的推理行为[5][6]。",
       "此外，Qwen2.5-Math 等开源数学模型也表明，面向数学任务的持续训练、"
       "自我改进和后训练策略能够显著提升模型的数学推理能力[8]。"),
      ("与传统 PPO 相比", "[4]"),
      ("减少了显式价值模型的需求", "[5]")]),
    # ---- 2.3 基座模型与参数高效微调 ----
    ("本文选择 Qwen2.5-1.5B-Instruct 作为主实验模型",
     [("作为主实验模型", "[7]")]),
    ("本文所有 SFT 和 GRPO 实验均采用 LoRA 方式进行参数高效微调",
     [("参数高效微调", "[3]，具体实现基于 Hugging Face PEFT 框架[10]")]),
    # ---- 2.4 数据集构造与答案抽取 ----
    ("本文最终使用 GSM8K 作为核心任务数据集",
     [("最终答案一般为单个数值", "[2]")]),
    ("在 SFT 数据方面，本文没有直接采用",
     [("元组答案和选择题答案的数据集", "[13]"),
      ("作为 CoT-SFT 数据来源", "[12]")]),
    ("GRPO 阶段使用 openai/gsm8k train split",
     [("中的 800 条样本", "[11]")]),
    ("最终评测使用完整的 openai/gsm8k test split",
     [("共 1319 条样本", "[11]")]),
    # ---- 2.5 CoT-SFT ----
    ("CoT-SFT 的目标是让小模型学习",
     [("规范化答案收束”的输出模式", "[1]")]),
    # ---- 2.6 GRPO ----
    ("GRPO（Group Relative Policy Optimization）是一种面向大语言模型后训练的策略优化方法",
     [("后训练的策略优化方法", "[5]"),
      ("与传统 PPO 使用价值模型估计优势不同", "[4]"),
      ("进行相对归一化", "[5]")]),
    ("训练时使用类似 PPO 的裁剪目标限制策略更新幅度",
     [("不要过度偏离参考模型。", "GRPO 训练基于 TRL 的 GRPOTrainer 实现[9]。")]),
    # ---- 2.7 奖励函数设计 ----
    ("本文设计三种规则奖励函数进行消融实验",
     [("不依赖人工偏好标注", "，与 RLVR 思路一致[5][6]")]),
    # ---- 3.1 实验平台与实验内容 ----
    ("实验在单卡 A800 80GB GPU 环境中完成",
     [("PyTorch、Transformers、TRL、PEFT 和 datasets",
       "，其中 GRPO 训练主要依赖 TRL，LoRA 微调主要依赖 PEFT[9][10]"),
      ("共 1319 道题", "[11]")]),
    # ---- 3.10 数据源选择与预实验反思 ----
    ("后训练数据源的选择会直接影响模型最终行为",
     [("证明题、选择题或表达式答案", "，例如 OpenR1-Math 这类通用数学数据[13]")]),
    ("它说明本文并不是机械套用 DeepSeek-R1 风格流程",
     [("机械套用 DeepSeek-R1 风格流程", "[6]")]),
]


def _segments(marker):
    """Split a marker into (text, is_citation) parts; citation parts are runs
    of one or more [n] groups that should become superscript."""
    out = []
    for part in re.split(r"((?:\[\d+\])+)", marker):
        if part:
            out.append((part, bool(re.fullmatch(r"(?:\[\d+\])+", part))))
    return out


def _make_run(text, base_rpr, sup):
    """Build a w:r run cloning base_rpr's font; superscript if sup=True."""
    r = OxmlElement("w:r")
    rpr = copy.deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")
    for va in rpr.findall(qn("w:vertAlign")):
        rpr.remove(va)
    if sup:
        va = OxmlElement("w:vertAlign"); va.set(qn("w:val"), "superscript"); rpr.append(va)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r


def ins(p, anchor, marker):
    """Insert `marker` right after `anchor`, splitting it into normal-text runs
    and SUPERSCRIPT runs for the [n] citations. Run-level => existing runs,
    fonts and inline Word equations are preserved."""
    full = p.text
    pos = full.find(anchor)
    if pos < 0:
        return False
    end = pos + len(anchor)
    acc = 0
    for run in p.runs:
        rl = len(run.text)
        if acc + rl >= end:
            local = end - acc
            base = run._element.find(qn("w:rPr"))
            tail = run.text[local:]
            run.text = run.text[:local]
            prev = run._element
            for text, is_cite in _segments(marker):
                el = _make_run(text, base, is_cite)
                prev.addnext(el); prev = el
            if tail:
                prev.addnext(_make_run(tail, base, False))
            return True
        acc += rl
    if p.runs:                        # offset at the very end of the paragraph
        run = p.runs[-1]; base = run._element.find(qn("w:rPr")); prev = run._element
        for text, is_cite in _segments(marker):
            el = _make_run(text, base, is_cite)
            prev.addnext(el); prev = el
        return True
    return False


def clone_para_after(ref, text):
    """Add a paragraph right after `ref`, cloning its style and run font."""
    newp = OxmlElement("w:p")
    ref._p.addnext(newp)
    np = Paragraph(newp, ref._parent)
    rp = ref._p.find(qn("w:pPr"))
    if rp is not None:
        np._p.insert(0, copy.deepcopy(rp))
    r = np.add_run(text)
    src = ref.runs[0]._element.find(qn("w:rPr")) if ref.runs else None
    if src is not None:
        r._element.insert(0, copy.deepcopy(src))
    return np


# missing reference entries to append after [11] (the in-text [12]/[13] need them)
EXTRA_REFS = [
    "[12] CAMEL-AI. gsm8k_distilled: GSM8K distilled chain-of-thought dataset. "
    "Hugging Face Dataset.",
    "[13] Open-R1. OpenR1-Math-220k dataset. Hugging Face Dataset.",
]


def main():
    shutil.copy(DOCX, DOCX + ".bak")
    d = docx.Document(DOCX)
    ok = bad = 0
    for locator, inserts in OPS:
        para = next((p for p in d.paragraphs if locator in p.text), None)
        if para is None:
            print(f"[MISS paragraph] {locator[:24]}")
            bad += len(inserts)
            continue
        for anchor, marker in inserts:
            if ins(para, anchor, marker):
                ok += 1
            else:
                bad += 1
                print(f"[MISS anchor] {anchor[:22]} | in: {locator[:18]}")

    # complete the reference list ([12]/[13] are cited but were missing)
    have = {p.text.strip()[:4] for p in d.paragraphs}
    ref11 = next((p for p in d.paragraphs if p.text.strip().startswith("[11]")), None)
    added = 0
    if ref11 is not None:
        anchor = ref11
        for entry in EXTRA_REFS:
            if entry[:4] not in have:
                anchor = clone_para_after(anchor, entry)
                added += 1
    d.save(DOCX)
    print(f"inserted {ok} citation markers ({bad} failed); appended {added} references. "
          f"backup: {DOCX}.bak")


if __name__ == "__main__":
    main()
