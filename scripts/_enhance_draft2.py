#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Fix formatting + insert architecture & unused figures into the latest draft."""
import re, shutil, sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
try: sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception: pass

SRC = "docs/深度学习大作业_邹研泽.docx"
DST = "docs/深度学习大作业_邹研泽_修订.docx"
FIG = "results/figures"
shutil.copy(SRC, DST)
d = docx.Document(DST)


def find(sub, style=None):
    s = sub.strip()
    for p in d.paragraphs:
        if s in p.text and (style is None or p.style.name == style):
            return p
    raise SystemExit("anchor not found: " + sub)


def has_img(p):
    return "graphicData" in p._p.xml or "pic:pic" in p._p.xml


def para_replace(p, fn):
    """Rewrite whole-paragraph text (robust to numbers split across runs).

    Skips paragraphs holding equations (OMML) or images so they are never
    disturbed. Flattens intra-paragraph run formatting, which is fine for
    plain captions / reference sentences.
    """
    if "oMath" in p._p.xml or has_img(p):
        return
    full = p.text
    nt = fn(full)
    if nt != full and p.runs:
        p.runs[0].text = nt
        for r in p.runs[1:]:
            r.text = ""


def set_style(np, style):
    try: np.style = d.styles[style]
    except Exception:
        try: np.style = style
        except Exception: pass


def para_before(anchor, text="", style="Normal", center=False, size=None):
    p = OxmlElement('w:p'); anchor._p.addprevious(p)
    np = Paragraph(p, anchor._parent); set_style(np, style)
    if text:
        r = np.add_run(text)
        if size: r.font.size = Pt(size)
    if center: np.alignment = 1
    return np


def img_before(anchor, img, width):
    p = OxmlElement('w:p'); anchor._p.addprevious(p)
    np = Paragraph(p, anchor._parent); np.alignment = 1
    np.add_run().add_picture(str(Path(FIG) / img), width=Inches(width))
    return np


import copy as _copy

def clone_para_before(anchor, text, ref):
    """Insert a paragraph before `anchor`, cloning `ref`'s full paragraph
    properties (style + line spacing + indent) and run font (incl. CJK font) so
    the new text matches the surrounding text in font, size and line spacing."""
    p = OxmlElement('w:p'); anchor._p.addprevious(p)
    np = Paragraph(p, anchor._parent)
    ref_pPr = ref._p.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = _copy.deepcopy(ref_pPr)          # style + spacing + indent
        for npr in new_pPr.findall(qn('w:numPr')):  # never inherit list numbering
            new_pPr.remove(npr)
        np._p.insert(0, new_pPr)
    else:
        np.style = ref.style
    r = np.add_run(text)
    if ref.runs:
        src = ref.runs[0]._element.find(qn('w:rPr'))
        if src is not None:
            r._element.insert(0, _copy.deepcopy(src))
    return np


from docx.oxml.ns import qn
import copy

# the document-level (body) sectPr governs the un-broken pages; clone its page
# size + margins into every inserted break so all sections have identical
# geometry (same text width => same column width everywhere).
_BODY_SECT = d.element.body.findall(qn('w:sectPr'))[-1]
_BODY_SPACE = (_BODY_SECT.find(qn('w:cols')).get(qn('w:space'))
               if _BODY_SECT.find(qn('w:cols')) is not None else "346")


def add_sectpr(paragraph, ncols, space=None):
    """Attach a *continuous* section break to paragraph: the section ENDING at
    this paragraph gets `ncols` columns. Page size / margins are cloned from the
    body section so column width matches the rest of the document."""
    if space is None:
        space = _BODY_SPACE
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:sectPr')):
        pPr.remove(old)
    sectPr = OxmlElement('w:sectPr')
    typ = OxmlElement('w:type'); typ.set(qn('w:val'), 'continuous'); sectPr.append(typ)
    for tag in ('w:pgSz', 'w:pgMar'):
        el = _BODY_SECT.find(qn(tag))
        if el is not None:
            sectPr.append(copy.deepcopy(el))
    cols = OxmlElement('w:cols'); cols.set(qn('w:num'), str(ncols)); cols.set(qn('w:space'), space)
    sectPr.append(cols)
    pPr.append(sectPr)


def find_cap(key):
    """Find a *caption* paragraph (one that starts with 图) containing `key`.

    Plain find() can match a body sentence that happens to quote the same
    phrase; captions always begin with 图, so anchor on that.
    """
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("图") and key in t:
            return p
    raise SystemExit("caption not found: " + key)


def _cap_index(cap):
    cs = list(d.paragraphs)
    return cs, next(i for i, p in enumerate(cs) if p._p is cap._p)


def replace_embedded(cap_key, png, width, new_caption=None, span=False):
    """Swap the image in the paragraph before caption `cap_key`; optionally
    rewrite the caption text and make the [image+caption] span both columns."""
    cap = find_cap(cap_key)
    cs, k = _cap_index(cap)
    imgp = cs[k - 1]
    for r in list(imgp.runs):
        r._element.getparent().remove(r._element)
    imgp.alignment = 1
    imgp.add_run().add_picture(str(Path(FIG) / png), width=Inches(width))
    if new_caption:
        cap.runs[0].text = new_caption
        for r in cap.runs[1:]:
            r.text = ""
    if span:
        add_sectpr(cs[k - 2], 2)
        add_sectpr(cap, 1)
    return cap


def delete_fig(cap_key):
    """Remove a figure's [image + caption] pair by caption keyword."""
    cap = find_cap(cap_key)
    cs, k = _cap_index(cap)
    img = cs[k - 1]
    cap._p.getparent().remove(cap._p)
    img._p.getparent().remove(img._p)


# ============ PART B: insert figures (temp numbers, fixed later) ============
# --- merge the 3 method diagrams into ONE combined (a)(b)(c) overview figure,
#     placed where 图1 (pipeline) was, spanning both columns ---
cap1 = find_cap("研究总体流程示意图")
caps = list(d.paragraphs)
ci = next(i for i, p in enumerate(caps) if p._p is cap1._p)
img_p1 = caps[ci - 1]          # the embedded (old, vertical) pipeline image
for r in list(img_p1.runs):    # drop the old picture
    r._element.getparent().remove(r._element)
img_p1.alignment = 1
img_p1.add_run().add_picture(str(Path(FIG) / "method_overview.png"), width=Inches(6.9))
cap1.runs[0].text = ("图 1 研究方法总览：(a) 总体实验流程与六组实验（E1–E5）；"
                     "(b) 数据构造与答案抽取（strict / lenient 双模式）；"
                     "(c) GRPO 组内相对优势机制。")
for r in cap1.runs[1:]:
    r.text = ""
add_sectpr(caps[ci - 2], 2)    # close the 2-column region before the figure
add_sectpr(cap1, 1)            # [image + caption] spans full text width

# the standalone data-pipeline (图2) & GRPO (图3) figures are now panels (b)/(c)
# of the combined figure -> remove their image + caption paragraphs.
for key in ("数据处理与答案抽取流程示意图", "GRPO训练机制示意图"):
    delete_fig(key)

# model architecture -> end of 2.2 (before 2.3)
a23 = find("2.3 数据集构造与答案抽取")
arch_lead = para_before(a23, "为更直观地展示基座模型结构与 LoRA 注入位置，模型整体架构如图 90 所示。"
            "LoRA 仅在自注意力与前馈网络的线性投影层上引入低秩增量，原始权重保持冻结，"
            "因此可训练参数量大幅降低。")
img_before(a23, "model_arch_h.png", 6.8)        # horizontal, full text width
arch_cap = para_before(a23, "图 90 Qwen2.5-1.5B 解码器结构与 LoRA 参数高效微调示意图。", center=True)
# carve out a full-width (1-column) region for the wide architecture figure
add_sectpr(arch_lead, 2)   # close the 2-column region just before the figure
add_sectpr(arch_cap, 1)    # the [image + caption] region spans full width

# radar -> end of 3.2 (before 3.3)
a33 = find("3.3 CoT-SFT 对格式规范性和准确率的影响")
para_before(a33, "为综合对比各方法，图 91 以雷达图形式展示六种方法在严格准确率、格式成功率、"
            "答案抽取率、输出简洁性与低过度推理五个维度上的表现（各维度按方法间最小最大值归一化）。"
            "可以看到，GRPO 在准确率维度领先但在简洁性与低过度推理维度较弱，SFT-GRPO 系列则在"
            "格式、抽取与简洁性维度更均衡。")
img_before(a33, "radar.png", 3.3)
para_before(a33, "图 91 各方法多维指标综合对比（雷达图，各维度已归一化）。", center=True)

# --- results figures: restyle + merge (numbers are remapped at the end) ---
# 图3 + 图4 -> one grouped bar (strict vs lenient); delete the lenient one first
delete_fig("上的宽松准确率")
replace_embedded("上的严格准确率", "accuracy_strict_lenient.png", 3.3,
                 "图 3 各方法在 GSM8K test 上的严格与宽松准确率对比。")
# format success + strict extraction (restyle, double bar)
replace_embedded("格式成功率与严格答案抽取", "format_extraction_rate.png", 3.3)
# average length + overthinking (now a dual-axis DOUBLE-BAR chart)
replace_embedded("平均输出长度与过度推理率", "length_overthinking.png", 3.3)

# 3.6: reward/loss/kl -> ONE combined (a)(b)(c) figure, spanning both columns
replace_embedded("各组训练奖励曲线", "train_curves_combined.png", 6.9,
                 "图 8 GRPO 训练过程：(a) 训练奖励；(b) 训练损失；(c) KL 散度"
                 "（淡线为原始值，粗线为 EMA 平滑）。", span=True)
# a sentence covering panels (b)(c) at the end of 3.6
a_len = find("3.7 输出长度与正确率关系分析")
para_before(a_len, "此外，训练损失与 KL 散度见图 8（b）（c）：损失整体下降后趋于小幅震荡，"
            "KL 散度始终维持在较小范围，说明策略更新幅度受控、训练过程稳定，未出现策略发散。")

# 3.7: the three length-analysis figures -> ONE combined (a)(b)(c), spanning
replace_embedded("输出长度与正确率关系分析图", "length_analysis_combined.png", 6.9,
                 "图 10 输出长度分析：(a) 分桶严格准确率；(b) 平均长度与严格准确率；"
                 "(c) 各方法输出长度分布。", span=True)
delete_fig("平均输出长度与严格准确率的关系")
# error-type composition (restyle)
replace_embedded("各方法错误类型构成", "error_type_distribution.png", 3.3)

# ============ PART C: add principle explanations + reflections ============
SFT_REF = find("strip_think")          # a 2.4 body paragraph (font reference)
GRPO_REF = find("num_generations")     # a 2.5 body paragraph (font reference)

# --- NEW section 2.1 (method principles). First shift existing 2.x -> 2.(x+1)
#     and the ablation sub-subsections 2.6.x -> 2.7.x (no body text references
#     section numbers, so this is safe). ---
def _bump_ch2():
    for p in d.paragraphs:
        st = p.style.name; t = p.text.strip()
        if st == "Heading 2":
            m = re.match(r"^2\.(\d+)\s*(.*)$", t)
            if m and p.runs:
                p.runs[0].text = f"2.{int(m.group(1)) + 1} {m.group(2)}"
                for r in p.runs[1:]: r.text = ""
        elif st == "Heading 3":
            m = re.match(r"^2\.6\.(\d+)\s*(.*)$", t)
            if m and p.runs:
                p.runs[0].text = f"2.7.{m.group(1)} {m.group(2)}"
                for r in p.runs[1:]: r.text = ""
_bump_ch2()

a_principle = find("任务定义与总体思路")     # now "2.2 任务定义..." -> insert before it
H2_REF_M = a_principle                        # Heading 2 style reference
H3_REF_M = find("正确性奖励")                 # "2.7.1 R1：正确性奖励" -> Heading 3 ref

clone_para_before(a_principle, "2.1 GRPO 与 SFT 方法原理", H2_REF_M)
clone_para_before(a_principle,
    "在展开具体实现之前，本节先对本实验涉及的两类核心后训练方法——强化学习方法 GRPO 与监督微调方法 SFT"
    "——的基本原理做简要梳理，作为后续实现与结果分析的基础。", GRPO_REF)

clone_para_before(a_principle, "2.1.1 GRPO（组相对策略优化）", H3_REF_M)
clone_para_before(a_principle,
    "GRPO（Group Relative Policy Optimization，组相对策略优化）是一种面向大语言模型的强化学习后训练方法。"
    "其机制可从策略梯度的角度理解：策略梯度以奖励为权重放大或抑制某条生成轨迹的概率，但原始奖励方差较大，"
    "直接使用会使训练不稳定，因此通常需要一个基线（baseline）做“相对好坏”的比较以降低方差。传统 PPO 使用一个"
    "与策略同规模的价值网络估计该基线，而 GRPO 的关键改进在于：对同一道题采样一组（G 个）回答，"
    "直接以该组奖励的均值作为基线、以标准差做归一化，从而得到每个回答的组内相对优势 A_i；在此基础上沿用 PPO 的"
    "裁剪式替代目标限制单步更新幅度，并对参考模型施加 KL 惩罚以防策略偏离过远：", GRPO_REF)
img_before(a_principle, "grpo_formula.png", 3.25)
clone_para_before(a_principle,
    "本人对该方法的体会是，GRPO 的“相对”二字才是关键：它并不追求绝对的奖励高低，而是不断让模型在“同一道题的"
    "一组尝试”中向更好的回答靠拢。这样既省去了价值网络的显存开销，又避免了人工偏好标注，"
    "因而非常契合本实验单卡、小模型的课程设计场景。", GRPO_REF)

clone_para_before(a_principle, "2.1.2 SFT（监督微调）", H3_REF_M)
clone_para_before(a_principle,
    "SFT（Supervised Fine-Tuning，监督微调）是后训练中最基础的方法，本质上是一种模仿学习。给定题目 x 与带有"
    "推理过程和最终答案的目标文本 y，SFT 以标准自回归语言建模损失训练模型，使其逐 token 拟合目标序列：", SFT_REF)
img_before(a_principle, "sft_formula.png", 2.7)
clone_para_before(a_principle,
    "训练时采用 teacher forcing，模型在每一步都以真实前文 token 为条件预测下一个 token，因此它学到的并非"
    "“如何推理”本身，而是“如何复现教师示范的推理轨迹”。本实验所用的思维链蒸馏，蒸的正是大模型解题时展开的中间"
    "步骤——把这种逐步求解的书写范式迁移到小模型上。", SFT_REF)
clone_para_before(a_principle,
    "需要说明的是，SFT 存在固有局限：模型至多逼近教师数据的分布，难以产生训练集中从未出现的新解法；并且训练阶段"
    "以真实前文为条件、推理阶段却以自身生成的前文为条件，存在暴露偏差。正因如此，本实验在 SFT 之后进一步引入以"
    "最终结果为奖励信号的 GRPO，以期突破单纯模仿的能力上限。", SFT_REF)

# Draft C -> new section 3.11 before the 4. 结论 heading
H2_REF = find("数据源选择与预实验反思")   # heading style reference (H2)
BODY_REF = SFT_REF
h4 = None
for p in d.paragraphs:
    s = p.text.strip().replace(" ", "")
    if s.endswith("结论") and len(s) <= 6:
        h4 = p; break
if h4 is None:
    raise SystemExit("conclusion heading not found")
clone_para_before(h4, "3.11 学习体会与实验心得", H2_REF)
clone_para_before(h4,
    "本实验的推进过程，对本人而言更像一次从“套用流程”到“理解机制”的学习过程。最初本人倾向于直接照搬 "
    "DeepSeek-R1 式的长思维链加强化学习范式，但预实验很快暴露问题：竞赛风格的长 CoT 数据与 GSM8K 的题目分布、"
    "答案格式严重错配，模型学到了不适合本任务的书写风格。这让本人体会到，后训练并非方法越“高级”越好，"
    "数据与目标任务的一致性往往比模型规模或方法新颖度更关键。", BODY_REF)
clone_para_before(h4,
    "在方法选型上，本实验之所以采用 SFT→GRPO 的两段式，是基于对二者作用边界的理解：SFT 擅长快速注入"
    "“规范输出格式”与“分步书写范式”，但只能模仿、难以超越教师；GRPO 则以最终答案正确性为信号，"
    "允许模型在自我采样中探索更优解。至于强化学习算法，相比需要额外价值网络的 PPO、以及依赖成对偏好数据的 DPO，"
    "GRPO 以“组内相对”的方式同时回避了价值网络的显存开销与人工偏好标注，在单卡小模型场景下最为务实，"
    "故成为本实验的选择。", BODY_REF)
clone_para_before(h4,
    "最令本人印象深刻的一点是“奖励即目标”：当奖励只奖励最终答案正确时，GRPO 组确实取得了最高准确率，"
    "却也产生了最长、最易过度推理的输出——模型忠实地优化了被给定的信号，而非本人真正想要的“既对又简洁”。"
    "这也正是本实验追加格式奖励与长度惩罚做消融的原因。尽管长度惩罚的实际效果有限，"
    "但这一系列结果让本人真切体会到：在强化学习后训练中，如何设计奖励，几乎等价于在定义“什么才是更好的回答”。",
    BODY_REF)

# ============ PART A: formatting fixes ============
# title (para 0): make it a real title
t0 = d.paragraphs[0]
set_style(t0, "Title")
if t0.runs:
    t0.runs[0].font.size = Pt(20); t0.runs[0].bold = True
t0.alignment = 1

# chapter headings: numbers + Heading 1 style
def set_heading(sub, new_text, level="Heading 1", exact_norm=False):
    for p in d.paragraphs:
        tx = p.text.strip()
        if (tx == sub) if exact_norm else (sub in tx):
            set_style(p, level)
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
            else:
                p.add_run(new_text)
            return True
    return False

set_heading("引言", "1. 引言", exact_norm=True)
set_heading("研究方法", "2. 研究方法", exact_norm=True)
set_heading("3. 实验与结果分析", "3. 实验与结果分析")
set_heading("4.结论", "4. 结论")
set_heading("4. 结论", "4. 结论")
set_heading("5. 参考文献", "5. 参考文献")

# some H1 (引言/研究方法/结论) carry an auto-number (numPr) on top of the manual
# "1." text -> renders "1.1.". Strip auto-numbering from every heading so only
# the manual section number shows.
for p in d.paragraphs:
    if p.style.name.startswith("Heading"):
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            npr = pPr.find(qn('w:numPr'))
            if npr is not None:
                pPr.remove(npr)

# duplicate 3.7 -> renumber later H2 and sub-subsections
def renum_h2(sub, new):
    for p in d.paragraphs:
        if p.style.name == "Heading 2" and sub in p.text:
            if p.runs:
                p.runs[0].text = re.sub(r"^\s*3\.\d+", new, p.text)
                for r in p.runs[1:]: r.text = ""
            return
renum_h2("典型错误类型分析", "3.8")
renum_h2("结果讨论", "3.9")
renum_h2("数据源选择与预实验反思", "3.10")
# sub-subsections 3.7.x -> 3.8.x (plain paragraphs; numbers may span runs)
for p in d.paragraphs:
    if re.match(r"^\s*3\.7\.\d", p.text):
        para_replace(p, lambda t: t.replace("3.7.", "3.8.", 1))

# image paragraph wrongly styled as Heading 2 -> Normal centered
for p in d.paragraphs:
    if has_img(p) and p.style.name == "Heading 2":
        set_style(p, "Normal"); p.alignment = 1

# typos + body refs to merged sub-panels (run level, before figure renumber).
# 图8 (reward) is now panel (a) of the training group; 图10 (scatter) is now
# panel (b) of the length-analysis group. Numbers are remapped afterwards.
TYPO = {"GSM8Ktest上": "GSM8K test 上", "GSM8Ktest": "GSM8K test",
        "图 8所示": "图 8（a）所示", "图 10以": "图 10（b）以"}
def _apply_typo(t):
    for a, b in TYPO.items():
        t = t.replace(a, b)
    return t
for p in d.paragraphs:
    if "GSM8Ktest" in p.text or "图 8所示" in p.text or "图 10以" in p.text:
        para_replace(p, _apply_typo)

# NOTE: the 3 contribution paragraphs already carry an auto-numbered list
# (numId=12 -> "1) 2) 3)"), consistent with the other lists in the document.
# Do NOT add a manual "（n）" prefix here -- that produced double numbering
# ("1) （1）"). Leaving the native list numbering intact.

# ============ renumber figures & tables + refs ============
fig_map = {}; tbl_map = {}; fi = ti = 0
for p in d.paragraphs:
    m = re.match(r"^\s*图\s*(\d+)", p.text)
    if m: fi += 1; fig_map[int(m.group(1))] = fi
    m = re.match(r"^\s*表\s*(\d+)", p.text)
    if m: ti += 1; tbl_map[int(m.group(1))] = ti

def sub(t):
    t = re.sub(r"图\s*(\d+)", lambda x: f"图 {fig_map.get(int(x.group(1)), x.group(1))}", t)
    t = re.sub(r"表\s*(\d+)", lambda x: f"表 {tbl_map.get(int(x.group(1)), x.group(1))}", t)
    return t

for p in d.paragraphs:
    if "图" in p.text or "表" in p.text:
        para_replace(p, sub)

d.save(DST)
print("Saved", DST)
print("fig_map:", fig_map)
print("tbl_map:", tbl_map)
